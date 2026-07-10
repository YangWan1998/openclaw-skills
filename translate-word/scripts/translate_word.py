#!/usr/bin/env python3
"""
Translate Word document (.docx) using provided translation mappings.

Usage:
    python translate_word.py <input.docx> <translations.json> <output.docx>

Arguments:
    input.docx         Input Word document path
    translations.json  JSON file with translation mappings: {"original": "translated", ...}
    output.docx        Output Word document path

注意：此脚本需要 Python 3.11+ 和 python-docx
如果当前环境不满足，会自动尝试使用 conda 的 py311 环境
"""

import sys
import subprocess
from pathlib import Path

# 检查 Python 版本，如果不满足则尝试切换环境
if sys.version_info < (3, 11):
    print(f"当前 Python {sys.version_info.major}.{sys.version_info.minor} 版本过低，需要 3.11+")
    print("尝试使用 conda py311 环境...")
    
    # 尝试使用 conda 环境重新运行
    try:
        result = subprocess.run(
            ["conda", "run", "-n", "py311", "python", __file__] + sys.argv[1:],
            capture_output=False,
            text=True
        )
        sys.exit(result.returncode)
    except Exception as e:
        print(f"切换环境失败: {e}")
        print("请手动运行: conda run -n py311 python translate_word.py")
        sys.exit(1)

import json
import argparse

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from copy import deepcopy
except ImportError:
    print("Error: python-docx not installed. Run: conda run -n py311 pip install python-docx")
    sys.exit(1)


def normalize_text(text: str) -> str:
    """Normalize text for fuzzy matching.
    
    Handles common encoding differences like:
    - Bullet symbols: · (U+00B7) vs • (U+2022) vs ● (U+25CF)
    - Quote symbols: " (U+0022) vs "" (U+201C/U+201D) vs '' (U+2018/U+2019)
    - Whitespace variations
    """
    if not text:
        return ""
    
    result = text
    
    # Replace various bullet symbols with a common placeholder
    bullets = ['\u00b7', '\u2022', '\u25cf', '\u25cb', '\u2013', '\u2014', '-']
    for bullet in bullets:
        result = result.replace(bullet, '\u00b7')
    
    # Replace various quote symbols with straight quotes
    result = result.replace('\u201c', '"').replace('\u201d', '"')  # " " -> "
    result = result.replace('\u2018', "'").replace('\u2019', "'")  # ' ' -> '
    
    # Normalize whitespace
    result = ' '.join(result.split())
    return result


def load_translations(translations_path: str) -> dict:
    """Load translation mappings from JSON file."""
    with open(translations_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # Support both flat dict and nested format
    if isinstance(data, dict):
        if "translations" in data:
            raw_translations = data["translations"]
        else:
            raw_translations = data
    else:
        return {}
    
    # Build normalized lookup table for fuzzy matching
    translations = {}
    normalized_lookup = {}
    
    for key, value in raw_translations.items():
        translations[key] = value
        normalized_key = normalize_text(key)
        if normalized_key and normalized_key not in normalized_lookup:
            normalized_lookup[normalized_key] = key  # Map normalized -> original key
    
    return translations, normalized_lookup


def find_translation(text: str, translations: dict, normalized_lookup: dict) -> str:
    """Find translation for text, supporting exact and fuzzy matching."""
    if not text or not text.strip():
        return None
    
    text = text.strip()
    
    # 1. Exact match
    if text in translations:
        return translations[text]
    
    # 2. Fuzzy match using normalized text
    normalized = normalize_text(text)
    if normalized in normalized_lookup:
        original_key = normalized_lookup[normalized]
        return translations.get(original_key)
    
    return None


def translate_paragraph(paragraph, translations: dict, normalized_lookup: dict, stats: dict):
    """Translate text in a paragraph while preserving formatting."""
    full_text = paragraph.text.strip()
    
    if not full_text:
        return
    
    # Check if entire paragraph has a translation (exact or fuzzy)
    translated = find_translation(full_text, translations, normalized_lookup)
    
    if translated and translated != full_text:
        # Clear the paragraph and rewrite
        paragraph.clear()
        run = paragraph.add_run(translated)
        
        # Try to preserve basic formatting from first run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            run.bold = first_run.bold
            run.italic = first_run.italic
            run.underline = first_run.underline
            
            # Preserve font size if available
            if first_run.font.size:
                run.font.size = first_run.font.size
            
            # Preserve color if available
            if first_run.font.color.rgb:
                run.font.color.rgb = first_run.font.color.rgb
        
        stats["translated_paragraphs"] += 1
        return
    
    # If no exact match, try to translate individual runs
    for run in paragraph.runs:
        original = run.text
        translated = find_translation(original, translations, normalized_lookup)
        if translated and translated != original:
            run.text = translated
            stats["translated_runs"] += 1


def translate_table(table, translations: dict, normalized_lookup: dict, stats: dict):
    """Translate text in table cells."""
    for row in table.rows:
        for cell in row.cells:
            # Try to match the entire cell text first (for multi-paragraph cells)
            cell_text = cell.text.strip()
            translated = find_translation(cell_text, translations, normalized_lookup)
            
            if translated and translated != cell_text:
                # Replace all paragraphs in the cell with the translated text
                # Split translation by newlines to preserve paragraph structure
                translated_lines = translated.split('\n')
                
                # Clear all existing paragraphs except the first one
                first_para = cell.paragraphs[0]
                first_para.clear()
                first_para.add_run(translated_lines[0] if translated_lines else translated)
                
                # Remove extra paragraphs
                for para in cell.paragraphs[1:]:
                    p_element = para._element
                    p_element.getparent().remove(p_element)
                
                # Add remaining lines as new paragraphs
                for line in translated_lines[1:]:
                    new_para = cell.add_paragraph(line)
                
                stats["translated_paragraphs"] += 1
            else:
                # Fall back to paragraph-by-paragraph translation
                for paragraph in cell.paragraphs:
                    translate_paragraph(paragraph, translations, normalized_lookup, stats)


def translate_word(input_path: str, translations: dict, normalized_lookup: dict, output_path: str):
    """
    Translate Word document using provided translation mappings.
    
    Args:
        input_path: Path to input .docx file
        translations: Dict mapping original text to translated text
        normalized_lookup: Dict mapping normalized text to original key
        output_path: Path for output .docx file
    
    Returns:
        dict: Translation statistics
    """
    doc = Document(input_path)
    
    stats = {
        "total_paragraphs": 0,
        "translated_paragraphs": 0,
        "translated_runs": 0,
        "translated_tables": 0,
        "skipped": 0
    }
    
    # Translate paragraphs
    for paragraph in doc.paragraphs:
        stats["total_paragraphs"] += 1
        translate_paragraph(paragraph, translations, normalized_lookup, stats)
    
    # Translate tables
    for table in doc.tables:
        translate_table(table, translations, normalized_lookup, stats)
        stats["translated_tables"] += 1
    
    # Save document
    doc.save(output_path)
    
    return stats


def main():
    parser = argparse.ArgumentParser(description="Translate Word document while preserving structure")
    parser.add_argument("input_docx", help="Input Word document (.docx)")
    parser.add_argument("translations_json", help="JSON file with translations")
    parser.add_argument("output_docx", nargs="?", help="Output Word document (.docx). If omitted and --organize is used, will be placed in organized folder")
    parser.add_argument("--organize", "-O", action="store_true",
                        help="Organize output into a folder named after the input file (without extension)")
    
    args = parser.parse_args()
    
    input_path = Path(args.input_docx)
    
    # Determine output path
    if args.organize:
        folder_name = input_path.stem
        output_dir = input_path.parent / folder_name
        output_dir.mkdir(exist_ok=True)
        output_docx = args.output_docx or str(output_dir / f"{folder_name}_EN.docx")
    else:
        if not args.output_docx:
            # Default: append _EN before extension
            output_docx = str(input_path.parent / f"{input_path.stem}_EN.docx")
        else:
            output_docx = args.output_docx
    
    # Load translations
    translations, normalized_lookup = load_translations(args.translations_json)
    print(f"Loaded {len(translations)} translation mappings")
    print(f"Normalized lookup table: {len(normalized_lookup)} entries")
    
    # Perform translation
    stats = translate_word(args.input_docx, translations, normalized_lookup, output_docx)
    
    print(f"\nTranslation complete!")
    print(f"Total paragraphs: {stats['total_paragraphs']}")
    print(f"Translated paragraphs: {stats['translated_paragraphs']}")
    print(f"Translated runs: {stats['translated_runs']}")
    print(f"Translated tables: {stats['translated_tables']}")
    print(f"Output: {output_docx}")


if __name__ == "__main__":
    main()
