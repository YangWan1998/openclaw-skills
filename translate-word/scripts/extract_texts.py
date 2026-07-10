#!/usr/bin/env python3
"""
Extract all text content from a Word document (.docx) with structure preserved.

Usage:
    python extract_texts.py <input.docx> [--output <texts.json>]

Output JSON format:
{
  "metadata": {
    "title": "...",
    "paragraphs_count": 42,
    "tables_count": 3,
    "images_count": 5
  },
  "paragraphs": [...],
  "tables": [...],
  "texts_to_translate": [...]
}

注意：此脚本需要 Python 3.11+ 和 python-docx
如果当前环境不满足，会自动尝试使用 conda 的 py311 环境
"""

import sys
import subprocess
from pathlib import Path

# 检查 Python 版本，如果不满足则尝试切换环境
if sys.version_info < (3, 11):
    print(f"当前 Python {sys.version_info.major}.{sys.version_info.minor} 版本过低，需要 3.11+")
    print("尝试使用 conda py311 环境...")
    
    # 尝试使用 conda 环境重新运行
    try:
        result = subprocess.run(
            ["conda", "run", "-n", "py311", "python", __file__] + sys.argv[1:],
            capture_output=False,
            text=True
        )
        sys.exit(result.returncode)
    except Exception as e:
        print(f"切换环境失败: {e}")
        print("请手动运行: conda run -n py311 python extract_texts.py")
        sys.exit(1)

import json
import argparse

try:
    from docx import Document
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("Error: python-docx not installed. Run: conda run -n py311 pip install python-docx")
    sys.exit(1)


def get_paragraph_type(paragraph):
    """Determine paragraph type based on style name."""
    style_name = paragraph.style.name.lower() if paragraph.style else ""
    
    if "heading 1" in style_name or "标题 1" in style_name:
        return "heading1"
    elif "heading 2" in style_name or "标题 2" in style_name:
        return "heading2"
    elif "heading 3" in style_name or "标题 3" in style_name:
        return "heading3"
    elif "heading" in style_name or "标题" in style_name:
        return "heading"
    elif "caption" in style_name or "题注" in style_name:
        return "caption"
    elif "table" in style_name:
        return "table_text"
    else:
        return "body"


def get_alignment(paragraph):
    """Get paragraph alignment."""
    alignment = paragraph.alignment
    if alignment == WD_PARAGRAPH_ALIGNMENT.LEFT:
        return "left"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.CENTER:
        return "center"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.RIGHT:
        return "right"
    elif alignment == WD_PARAGRAPH_ALIGNMENT.JUSTIFY:
        return "justify"
    else:
        return "left"


def extract_paragraphs(doc):
    """Extract all paragraphs with their structure."""
    paragraphs = []
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        # Check formatting
        is_bold = any(run.bold for run in para.runs if run.bold is not None)
        is_italic = any(run.italic for run in para.runs if run.italic is not None)
        
        para_info = {
            "id": f"p{idx}",
            "type": get_paragraph_type(para),
            "text": text,
            "style": para.style.name if para.style else "Normal",
            "is_bold": is_bold,
            "is_italic": is_italic,
            "alignment": get_alignment(para)
        }
        paragraphs.append(para_info)
    
    return paragraphs


def extract_tables(doc):
    """Extract all tables with their cell contents."""
    tables = []
    
    for tidx, table in enumerate(doc.tables):
        cells = []
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append({
                        "row": row_idx,
                        "col": col_idx,
                        "text": cell_text
                    })
        
        if cells:
            tables.append({
                "id": f"t{tidx}",
                "rows": len(table.rows),
                "cols": len(table.rows[0].cells) if table.rows else 0,
                "cells": cells
            })
    
    return tables


def count_images(doc):
    """Count images in the document."""
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            image_count += 1
    return image_count


def extract_texts(input_path: str, output_path: str = None):
    """Extract all text content from Word document."""
    doc = Document(input_path)
    
    # Extract structure
    paragraphs = extract_paragraphs(doc)
    tables = extract_tables(doc)
    image_count = count_images(doc)
    
    # Collect all unique texts for translation
    all_texts = set()
    
    for para in paragraphs:
        all_texts.add(para["text"])
    
    for table in tables:
        for cell in table["cells"]:
            all_texts.add(cell["text"])
    
    result = {
        "metadata": {
            "title": doc.core_properties.title or "",
            "paragraphs_count": len(paragraphs),
            "tables_count": len(tables),
            "images_count": image_count
        },
        "paragraphs": paragraphs,
        "tables": tables,
        "texts_to_translate": sorted(list(all_texts))
    }
    
    if output_path:
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, ensure_ascii=False, indent=2)
        print(f"Extracted {len(paragraphs)} paragraphs, {len(tables)} tables, {image_count} images")
        print(f"Total unique texts: {len(all_texts)}")
        print(f"Saved to: {output_path}")
    
    return result


def main():
    parser = argparse.ArgumentParser(description="Extract text from Word document")
    parser.add_argument("input_docx", help="Input Word document (.docx)")
    parser.add_argument("--output", "-o", help="Output JSON file (optional)")
    parser.add_argument("--organize", "-O", action="store_true", 
                        help="Organize outputs into a folder named after the input file (without extension)")
    
    args = parser.parse_args()
    
    input_path = Path(args.input_docx)
    
    if args.organize:
        # Create organized folder: A.docx -> A/A.extracted.json
        folder_name = input_path.stem
        output_dir = input_path.parent / folder_name
        output_dir.mkdir(exist_ok=True)
        output = args.output or str(output_dir / f"{folder_name}.extracted.json")
    else:
        output = args.output or str(input_path.with_suffix(".extracted.json"))
    
    extract_texts(args.input_docx, output)


if __name__ == "__main__":
    main()
